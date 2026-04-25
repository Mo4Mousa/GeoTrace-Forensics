import html
import json
import subprocess
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
ASSETS_DIR = OUTPUT_DIR / "project_doc_assets"
REPORT_HTML = OUTPUT_DIR / "GeoTrace_Forensics_Project_Report.html"
REPORT_DOCX = OUTPUT_DIR / "GeoTrace_Forensics_Project_Report.docx"
REPORT_DOCX_EN = OUTPUT_DIR / "GeoTrace_Forensics_Project_Report_English.docx"


CASE_JSON_PATH = BASE_DIR / "output" / "case_4" / "case_4_export.json"
LOGO_PATH = BASE_DIR / "assets" / "project_logo.png"


def load_text(relative_path):
    return (BASE_DIR / relative_path).read_text(encoding="utf-8")


def load_json(path):
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dirs():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def font_candidates():
    return {
        "sans": [
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Supplemental/Tahoma.ttf",
            "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        ],
        "mono": [
            "/System/Library/Fonts/SFNSMono.ttf",
            "/System/Library/Fonts/Supplemental/Courier New.ttf",
            "/System/Library/Fonts/Supplemental/Menlo.ttc",
            "/Library/Fonts/MesloLGS NF Regular.ttf",
        ],
    }


def load_font(kind, size, bold=False):
    candidates = font_candidates()[kind][:]
    if bold and kind == "sans":
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/System/Library/Fonts/Supplemental/Tahoma Bold.ttf",
        ] + candidates

    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def fit_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""

    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines or [text]


def save_card(title, lines, output_path, width=1480, height=900):
    image = Image.new("RGB", (width, height), "#f5efe3")
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=28, fill="#fcfaf5", outline="#d7cbb7", width=3)
    draw.rounded_rectangle((60, 60, width - 60, 150), radius=20, fill="#264653")
    draw.ellipse((86, 89, 104, 107), fill="#ff6b6b")
    draw.ellipse((116, 89, 134, 107), fill="#ffd166")
    draw.ellipse((146, 89, 164, 107), fill="#06d6a0")

    title_font = load_font("sans", 34, bold=True)
    mono_font = load_font("mono", 26)
    draw.text((200, 86), title, font=title_font, fill="white")

    y = 195
    for line in lines:
        wrapped = textwrap.wrap(line, width=88, replace_whitespace=False) or [""]
        for item in wrapped:
            draw.text((92, y), item, font=mono_font, fill="#22333b")
            y += 38
        y += 8

    image.save(output_path)


def save_architecture_diagram(output_path):
    width, height = 1600, 950
    image = Image.new("RGB", (width, height), "#f7f3eb")
    draw = ImageDraw.Draw(image)
    title_font = load_font("sans", 38, bold=True)
    box_title_font = load_font("sans", 26, bold=True)
    body_font = load_font("sans", 22)

    draw.text((70, 48), "GeoTrace Forensics Workflow", font=title_font, fill="#1f2937")

    boxes = [
        ((90, 160, 460, 330), "#2a9d8f", "1. Evidence Intake", ["Create case", "Import images", "Store SHA-256"]),
        ((620, 160, 990, 330), "#e9c46a", "2. Metadata Extraction", ["Read EXIF", "Decode GPS", "Normalize raw tags"]),
        ((1120, 160, 1490, 330), "#f4a261", "3. Forensic Analysis", ["Detect anomalies", "Timeline ordering", "Duplicate correlation"]),
        ((210, 520, 580, 690), "#8ecae6", "4. Persistence Layer", ["SQLite cases", "Images", "Metadata + anomalies"]),
        ((740, 520, 1110, 690), "#90be6d", "5. Investigator Review", ["GUI table", "Raw EXIF view", "Integrity status"]),
        ((1240, 520, 1510, 690), "#e76f51", "6. Outputs", ["HTML map", "PDF report", "Excel + JSON export"]),
    ]

    for x1, y1, x2, y2 in [box[0] for box in boxes]:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#ffffff", outline="#d0c4b1", width=3)

    for coords, accent, title, items in boxes:
        x1, y1, x2, y2 = coords
        draw.rounded_rectangle((x1, y1, x2, y1 + 54), radius=24, fill=accent)
        draw.rectangle((x1, y1 + 28, x2, y1 + 54), fill=accent)
        draw.text((x1 + 20, y1 + 12), title, font=box_title_font, fill="#ffffff")

        item_y = y1 + 78
        for item in items:
            draw.ellipse((x1 + 24, item_y + 8, x1 + 36, item_y + 20), fill=accent)
            draw.text((x1 + 52, item_y), item, font=body_font, fill="#25303b")
            item_y += 40

    arrow_color = "#355070"
    arrows = [
        (460, 245, 620, 245),
        (990, 245, 1120, 245),
        (1300, 330, 1300, 520),
        (930, 330, 930, 520),
        (400, 330, 400, 520),
        (580, 605, 740, 605),
        (1110, 605, 1240, 605),
    ]
    for x1, y1, x2, y2 in arrows:
        draw.line((x1, y1, x2, y2), fill=arrow_color, width=8)
        if x2 > x1:
            draw.polygon([(x2, y2), (x2 - 26, y2 - 14), (x2 - 26, y2 + 14)], fill=arrow_color)
        elif y2 > y1:
            draw.polygon([(x2, y2), (x2 - 14, y2 - 26), (x2 + 14, y2 - 26)], fill=arrow_color)

    image.save(output_path)


def build_case_summary_lines(case_data):
    case_info = case_data["case"]
    image_info = case_data["images"][0]
    summary = case_data["timeline"]["summary"]
    anomalies = image_info.get("anomalies", [])

    return [
        f"Case ID: {case_info['case_id']}",
        f"Case Name: {case_info['case_name']}",
        f"Investigator: {case_info['investigator_name']}",
        f"Created At: {case_info['created_at']}",
        "",
        f"Image: {image_info['file_name']}",
        f"Camera: {image_info.get('camera_make') or 'N/A'} {image_info.get('camera_model') or ''}".strip(),
        f"Date Taken: {image_info.get('date_taken') or 'Unavailable'}",
        f"Integrity: {image_info.get('integrity_status') or 'Unknown'}",
        f"SHA-256: {image_info['sha256_hash']}",
        "",
        f"Images With GPS: {summary.get('images_with_gps', 0)}",
        f"Total Anomalies: {summary.get('total_anomalies', 0)}",
        "Detected Findings:",
        *[f"- {item['anomaly_type']} ({item['severity']}): {item['description']}" for item in anomalies],
    ]


def create_assets(case_data):
    tree_lines = [
        "GeoTrace_Forensics/",
        "  main.py",
        "  modules/",
        "    exif_extractor.py",
        "    gps_decoder.py",
        "    anomaly_detector.py",
        "    timeline_generator.py",
        "    map_generator.py",
        "    report_generator.py",
        "    export_manager.py",
        "    db_manager.py",
        "  ui/",
        "    main_window.py",
        "    style.qss",
        "  database/",
        "    schema.sql",
        "  output/",
        "    case_4/",
        "      case_4_export.json",
        "      case_4_export.xlsx",
        "      case_4_forensic_report.pdf",
    ]
    save_card("Project Structure Snapshot", tree_lines, ASSETS_DIR / "project_tree.png", height=980)

    case_summary_lines = build_case_summary_lines(case_data)
    save_card("Case 4 Findings Snapshot", case_summary_lines, ASSETS_DIR / "case4_summary.png", height=960)

    json_lines = CASE_JSON_PATH.read_text(encoding="utf-8").splitlines()[:34]
    save_card("JSON Export Snapshot", json_lines, ASSETS_DIR / "json_snapshot.png", height=1080)

    main_window_lines = load_text("ui/main_window.py").splitlines()[26:76]
    save_card("UI Controller Code Snapshot", main_window_lines, ASSETS_DIR / "ui_code.png", height=1160)

    exif_lines = load_text("modules/exif_extractor.py").splitlines()[0:68]
    save_card("EXIF Extraction Code Snapshot", exif_lines, ASSETS_DIR / "exif_code.png", height=1220)

    anomaly_lines = load_text("modules/anomaly_detector.py").splitlines()[0:78]
    save_card("Anomaly Detection Code Snapshot", anomaly_lines, ASSETS_DIR / "anomaly_code.png", height=1260)

    timeline_lines = load_text("modules/timeline_generator.py").splitlines()[0:90]
    save_card("Timeline Analysis Code Snapshot", timeline_lines, ASSETS_DIR / "timeline_code.png", height=1320)

    save_architecture_diagram(ASSETS_DIR / "architecture.png")


def snippet(relative_path, start_line, end_line):
    lines = load_text(relative_path).splitlines()
    selected = lines[start_line - 1:end_line]
    return "\n".join(selected)


def build_html(case_data):
    case_info = case_data["case"]
    image_info = case_data["images"][0]
    summary = case_data["timeline"]["summary"]
    anomaly_rows = image_info.get("anomalies", [])

    feature_rows = [
        ("EXIF Extraction", "Reading camera, time, software, and raw metadata from each image."),
        ("GPS Decoding", "Converting EXIF latitude/longitude to investigator-friendly decimal coordinates."),
        ("Timeline Analysis", "Sorting captures chronologically and measuring time/distance gaps."),
        ("Integrity Verification", "Using SHA-256 hashing to confirm that evidence has not changed."),
        ("Anomaly Detection", "Flagging missing GPS, suspicious timestamps, editing software, and modified files."),
        ("Forensic Outputs", "Generating PDF, Excel, JSON, and interactive HTML map artifacts."),
    ]

    module_rows = [
        ("main.py", "Application entry point and stylesheet loading."),
        ("ui/main_window.py", "Desktop interface for case creation, evidence intake, and artifact generation."),
        ("modules/exif_extractor.py", "Metadata extraction with exifread plus Pillow fallback."),
        ("modules/gps_decoder.py", "GPS normalization, validation, and coordinate formatting."),
        ("modules/anomaly_detector.py", "Forensic rules for suspicious metadata conditions."),
        ("modules/timeline_generator.py", "Chronological ordering and correlation analysis."),
        ("modules/map_generator.py", "Interactive folium map generation and timestamp playback."),
        ("modules/report_generator.py", "PDF report creation with case summary and evidence chain."),
        ("modules/export_manager.py", "Excel and JSON export packaging."),
        ("modules/db_manager.py", "SQLite persistence for cases, images, metadata, and anomalies."),
    ]

    anomalies_html = "".join(
        f"<tr><td>{html.escape(item['anomaly_type'])}</td><td>{html.escape(item['severity'])}</td><td>{html.escape(item['description'])}</td></tr>"
        for item in anomaly_rows
    )

    features_html = "".join(
        f"<tr><td>{html.escape(name)}</td><td>{html.escape(desc)}</td></tr>"
        for name, desc in feature_rows
    )
    modules_html = "".join(
        f"<tr><td>{html.escape(name)}</td><td>{html.escape(desc)}</td></tr>"
        for name, desc in module_rows
    )

    exif_snippet = html.escape(snippet("modules/exif_extractor.py", 41, 90))
    anomaly_snippet = html.escape(snippet("modules/anomaly_detector.py", 28, 91))
    timeline_snippet = html.escape(snippet("modules/timeline_generator.py", 33, 118))

    html_body = f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
  <meta charset="utf-8">
  <title>GeoTrace Forensics Project Report</title>
  <style>
    body {{
      font-family: Arial, Tahoma, sans-serif;
      margin: 36px;
      color: #1f2937;
      line-height: 1.7;
      background: #fffdfa;
    }}
    h1, h2, h3 {{
      color: #12324a;
      margin-bottom: 8px;
    }}
    h1 {{
      font-size: 28px;
      text-align: center;
    }}
    h2 {{
      font-size: 22px;
      border-bottom: 2px solid #d8c4a8;
      padding-bottom: 6px;
      margin-top: 28px;
    }}
    h3 {{
      font-size: 18px;
      margin-top: 18px;
    }}
    p, li, td, th {{
      font-size: 13.5px;
    }}
    .cover {{
      text-align: center;
      padding: 8px 0 18px;
    }}
    .cover img {{
      width: 420px;
      margin-bottom: 10px;
    }}
    .meta {{
      margin: 0 auto;
      width: 86%;
      border: 1px solid #d8c4a8;
      background: #f8f2e8;
      padding: 12px 18px;
      border-radius: 10px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 12px 0 18px;
    }}
    th {{
      background: #264653;
      color: white;
      padding: 9px;
      text-align: right;
    }}
    td {{
      border: 1px solid #d9d3ca;
      padding: 9px;
      vertical-align: top;
    }}
    .figure {{
      text-align: center;
      margin: 18px 0 24px;
    }}
    .figure img {{
      width: 92%;
      border: 1px solid #d9d3ca;
    }}
    .caption {{
      margin-top: 6px;
      color: #5c6770;
      font-size: 12px;
    }}
    pre {{
      direction: ltr;
      text-align: left;
      white-space: pre-wrap;
      background: #f5efe3;
      border: 1px solid #d9d3ca;
      padding: 14px;
      font-family: Menlo, Consolas, monospace;
      font-size: 11px;
      line-height: 1.45;
    }}
    .two-col {{
      width: 100%;
    }}
    .note {{
      background: #eef6f3;
      border-right: 5px solid #2a9d8f;
      padding: 10px 14px;
      margin: 12px 0 16px;
    }}
    .small {{
      font-size: 12px;
      color: #5c6770;
    }}
  </style>
</head>
<body>
  <div class="cover">
    <img src="{LOGO_PATH.as_posix()}" alt="GeoTrace Logo">
    <h1>توثيق احترافي لمشروع GeoTrace Forensics</h1>
    <p>Desktop Digital Forensics Tool for Image Metadata, Geolocation, Timeline, and Reporting</p>
    <div class="meta">
      <p><strong>اسم المشروع:</strong> GeoTrace Forensics</p>
      <p><strong>نوع المشروع:</strong> Desktop Forensic Investigation Application built with Python and PyQt5</p>
      <p><strong>تاريخ إعداد الوثيقة:</strong> 2026-04-25</p>
      <p><strong>الحالة العملية المستخدمة في التوثيق:</strong> Case #{case_info['case_id']} - {html.escape(case_info['case_name'])}</p>
    </div>
  </div>

  <h2>1. الملخص التنفيذي</h2>
  <p>
    هذا المشروع عبارة عن أداة تحقيق رقمي مكتبية تهدف إلى تحليل صور الأدلة الرقمية بطريقة منظمة واحترافية.
    يقوم النظام باستيراد الصور، استخراج بيانات EXIF، محاولة فك إحداثيات GPS، التحقق من سلامة الملفات باستخدام SHA-256،
    بناء تسلسل زمني للصور، اكتشاف الشذوذات الجنائية، ثم إخراج النتائج في صورة تقارير PDF وملفات Excel وJSON وخرائط HTML تفاعلية.
  </p>
  <p>
    من خلال مراجعة الكود والمخرجات الفعلية، يتضح أن المشروع لا يقتصر على عرض البيانات فقط، بل يطبق سلسلة معالجة جنائية متكاملة
    تبدأ من intake evidence وتنتهي بإخراج تقارير قابلة للمراجعة والتسليم.
  </p>

  <h2>2. هدف المشروع والمشكلة التي يحلها</h2>
  <p>
    في التحقيقات الرقمية، صور الهواتف والكاميرات قد تحتوي على معلومات حساسة تساعد المحقق في معرفة وقت الالتقاط، نوع الجهاز،
    البرنامج المستخدم في التعديل، وموقع الالتقاط الجغرافي. المشكلة أن هذه البيانات تكون متناثرة وغير جاهزة للتحليل اليدوي السريع.
    هنا يأتي دور GeoTrace Forensics، حيث يجمع هذه البيانات في واجهة واحدة، ويحوّلها إلى evidence package قابل للفحص القضائي والتحليلي.
  </p>

  <div class="figure">
    <img src="{(ASSETS_DIR / 'architecture.png').as_posix()}" alt="Architecture Diagram">
    <div class="caption">شكل 1: التسلسل العام لعمل النظام من إدخال الأدلة حتى إخراج النتائج.</div>
  </div>

  <h2>3. المزايا الأساسية للمشروع</h2>
  <table>
    <tr><th>الميزة</th><th>الوصف</th></tr>
    {features_html}
  </table>

  <h2>4. بنية المشروع والملفات الرئيسية</h2>
  <p>
    يعتمد المشروع على تقسيم واضح بين الواجهة الرسومية، وحدات التحليل الجنائي، قاعدة البيانات، ومجلد الإخراج الخاص بالقضايا.
    هذا التنظيم يسهل الصيانة والتوسعة ويجعل كل مسؤولية معزولة داخل ملف أو وحدة محددة.
  </p>
  <div class="figure">
    <img src="{(ASSETS_DIR / 'project_tree.png').as_posix()}" alt="Project Tree">
    <div class="caption">شكل 2: لقطة مبنية من هيكل المشروع الفعلي والملفات الأهم داخله.</div>
  </div>
  <table>
    <tr><th>الملف أو الوحدة</th><th>الدور داخل المشروع</th></tr>
    {modules_html}
  </table>

  <h2>5. شرح سير العمل داخل التطبيق</h2>
  <ol>
    <li>إنشاء قضية جديدة وإدخال اسم القضية واسم المحقق.</li>
    <li>استيراد الصور وربطها بالقضية الحالية.</li>
    <li>حساب SHA-256 لكل صورة وتسجيل حجم الملف ومساره.</li>
    <li>استخراج بيانات EXIF مثل الكاميرا ووقت الالتقاط والبرنامج المستخدم.</li>
    <li>فك وترميز إحداثيات GPS إن كانت موجودة داخل الصورة.</li>
    <li>تطبيق قواعد كشف الشذوذات الجنائية على كل صورة.</li>
    <li>تجميع النتائج في Timeline مرتب مع correlations عند وجود تقارب زمني ومكاني.</li>
    <li>إخراج النتائج في PDF وExcel وJSON، بالإضافة إلى خريطة تفاعلية HTML للصور التي تحتوي على GPS.</li>
  </ol>

  <h2>6. شرح الواجهة الرسومية</h2>
  <p>
    الواجهة الرئيسية مبنية باستخدام PyQt5 وتضم ثلاثة أجزاء أساسية: إدارة القضايا، شريط العمليات، ومنطقة استعراض النتائج.
    يمكن للمحقق إنشاء أو فتح قضية، ثم تنفيذ إجراءات مثل إضافة الصور، إنشاء التقرير، التحقق من سلامة الأدلة، والتصدير.
  </p>
  <div class="figure">
    <img src="{(ASSETS_DIR / 'ui_code.png').as_posix()}" alt="UI Code">
    <div class="caption">شكل 3: لقطة من كود الواجهة الرئيسية تظهر بناء العناصر الأساسية والتحكم في مسار التحقيق.</div>
  </div>

  <h2>7. قاعدة البيانات والتخزين الجنائي</h2>
  <p>
    يستخدم المشروع قاعدة SQLite محلية، وهي مناسبة لهذا النوع من الأدوات الأكاديمية والخفيفة. قاعدة البيانات تتكون من أربع جداول أساسية:
    <strong>cases</strong> لتخزين القضايا، <strong>images</strong> لتسجيل ملفات الأدلة، <strong>metadata</strong> لتخزين نتائج الاستخراج،
    و<strong>anomalies</strong> لتسجيل كل شذوذ أو ملاحظة جنائية مرتبطة بالصورة.
  </p>
  <div class="note">
    هذا التصميم يحقق الفصل بين evidence intake وmetadata enrichment وanomaly tracking، وهو قرار جيد من ناحية التنظيم وإمكانية التوسع.
  </div>

  <h2>8. التحليل العملي للحالة Case {case_info['case_id']}</h2>
  <p>
    تم استخدام الملف الناتج الحقيقي الموجود في <strong>output/case_4/case_4_export.json</strong> لشرح كيفية عمل الأداة عمليًا.
    تحتوي هذه الحالة على صورة واحدة هي <strong>{html.escape(image_info['file_name'])}</strong> تم التقاطها بواسطة
    <strong>{html.escape((image_info.get('camera_make') or '') + ' ' + (image_info.get('camera_model') or '')).strip()}</strong>.
  </p>

  <table>
    <tr><th>العنصر</th><th>القيمة</th></tr>
    <tr><td>Case ID</td><td>{case_info['case_id']}</td></tr>
    <tr><td>Case Name</td><td>{html.escape(case_info['case_name'])}</td></tr>
    <tr><td>Investigator</td><td>{html.escape(case_info['investigator_name'])}</td></tr>
    <tr><td>Image File</td><td>{html.escape(image_info['file_name'])}</td></tr>
    <tr><td>Capture Time</td><td>{html.escape(image_info.get('date_taken') or 'Unavailable')}</td></tr>
    <tr><td>GPS Status</td><td>{summary.get('images_with_gps', 0)} image(s) with GPS</td></tr>
    <tr><td>Integrity Status</td><td>{html.escape(image_info.get('integrity_status') or 'Unknown')}</td></tr>
    <tr><td>Total Anomalies</td><td>{summary.get('total_anomalies', 0)}</td></tr>
    <tr><td>SHA-256</td><td>{html.escape(image_info['sha256_hash'])}</td></tr>
  </table>

  <div class="figure">
    <img src="{(ASSETS_DIR / 'case4_summary.png').as_posix()}" alt="Case 4 Summary">
    <div class="caption">شكل 4: لقطة ملخصة لأهم نتائج Case 4 كما تم استخراجها من ملف التصدير الفعلي.</div>
  </div>

  <h3>8.1 أهم النتائج الجنائية في الحالة</h3>
  <p>
    أظهرت الحالة أن الملف تم التحقق من سلامته بنجاح لأن قيمة SHA-256 الحالية تطابقت مع القيمة المخزنة عند الإدخال.
    في المقابل، لم يتم العثور على إحداثيات GPS، وبالتالي لا يمكن تحديد الموقع الجغرافي مباشرة من هذه الصورة.
    كما لاحظ النظام أن وقت التعديل على مستوى نظام الملفات جاء بعد وقت الالتقاط المسجل داخل EXIF، ولهذا تم تسجيل شذوذ منخفض الخطورة.
  </p>

  <table>
    <tr><th>نوع الشذوذ</th><th>الخطورة</th><th>الوصف</th></tr>
    {anomalies_html}
  </table>

  <div class="figure">
    <img src="{(ASSETS_DIR / 'json_snapshot.png').as_posix()}" alt="JSON Snapshot">
    <div class="caption">شكل 5: لقطة من ملف JSON المولد فعليًا للحالة، وتظهر بيانات القضية والصورة والشذوذات.</div>
  </div>

  <h2>9. شرح الكود الأساسي</h2>
  <h3>9.1 استخراج EXIF</h3>
  <p>
    وحدة <strong>exif_extractor.py</strong> تقرأ الملف باستخدام Pillow لاستخراج البنية العامة للصورة،
    ثم تستخدم exifread لاستخراج الوسوم المتخصصة عند الإمكان. بعد ذلك يتم دمج البيانات وإرجاعها في بنية موحدة واحدة.
    هذه فكرة جيدة لأنها توفر fallback عملي إذا لم تنجح إحدى المكتبتين في قراءة نوع الصورة.
  </p>
  <div class="figure">
    <img src="{(ASSETS_DIR / 'exif_code.png').as_posix()}" alt="EXIF Code">
    <div class="caption">شكل 6: لقطة من كود استخراج EXIF ودمج الوسوم في Metadata موحدة.</div>
  </div>
  <pre>{exif_snippet}</pre>

  <h3>9.2 كشف الشذوذات</h3>
  <p>
    وحدة <strong>anomaly_detector.py</strong> تمثل قلب التحليل الجنائي في المشروع. فهي تراجع وجود EXIF، واكتمال GPS،
    وصيغة الوقت، واحتمالية وجود برامج تعديل، وكذلك مقارنة وقت التعديل على نظام الملفات مع وقت الالتقاط داخل الصورة.
    هذه القواعد تجعل الأداة مفيدة في الفحص الأولي والتحليل السريع.
  </p>
  <div class="figure">
    <img src="{(ASSETS_DIR / 'anomaly_code.png').as_posix()}" alt="Anomaly Code">
    <div class="caption">شكل 7: لقطة من كود القواعد الجنائية المستخدمة لاكتشاف المشكلات داخل Metadata.</div>
  </div>
  <pre>{anomaly_snippet}</pre>

  <h3>9.3 بناء الخط الزمني</h3>
  <p>
    بعد جمع النتائج، تقوم وحدة <strong>timeline_generator.py</strong> بترتيب الصور زمنيًا، وحساب الفروق الزمنية،
    وقياس المسافة بين الأحداث باستخدام معادلة Haversine عند توفر الإحداثيات. إذا كان الفرق الزمني والمكاني ضمن نافذة محددة،
    يتم إنشاء correlation يساعد في ربط الصور ببعضها ضمن سيناريو تحقيق واحد.
  </p>
  <div class="figure">
    <img src="{(ASSETS_DIR / 'timeline_code.png').as_posix()}" alt="Timeline Code">
    <div class="caption">شكل 8: لقطة من كود التحليل الزمني وحساب التقارب المكاني بين الصور.</div>
  </div>
  <pre>{timeline_snippet}</pre>

  <h2>10. المخرجات التي ينتجها النظام</h2>
  <p>يقوم المشروع بإنتاج أكثر من artifact يمكن للمحقق الاستفادة منه حسب مرحلة العمل:</p>
  <ul>
    <li><strong>PDF Report:</strong> تقرير منسق يلخص القضية، الأدلة، الخط الزمني، والشذوذات.</li>
    <li><strong>Excel Export:</strong> مناسب للفرز والتحليل اليدوي ومشاركة البيانات مع الفريق.</li>
    <li><strong>JSON Export:</strong> مناسب للتكامل مع أدوات أخرى أو التحليل البرمجي لاحقًا.</li>
    <li><strong>HTML Map:</strong> خريطة تفاعلية عند توفر صور تحتوي على إحداثيات GPS.</li>
  </ul>
  <p class="small">
    ملاحظة: الحالة العملية Case {case_info['case_id']} لا تحتوي على GPS، لذلك لم يتم توليد خريطة خاصة بها داخل المخرجات الحالية،
    رغم أن ميزة إنشاء الخرائط موجودة ومطبقة داخل المشروع.
  </p>

  <h2>11. طريقة تشغيل المشروع</h2>
  <pre>pip install -r requirements.txt
python main.py</pre>
  <p>
    عند التشغيل، يتم فتح الواجهة الرسومية مباشرة، ثم يمكن إنشاء قضية جديدة وإضافة الصور وبدء التحليل.
    يعتمد المشروع على المكتبات: PyQt5, Pillow, exifread, folium, reportlab, openpyxl, pandas.
  </p>

  <h2>12. نقاط القوة والملاحظات التطويرية</h2>
  <h3>نقاط القوة</h3>
  <ul>
    <li>تقسيم واضح للوحدات البرمجية وسهولة فهم مسؤولية كل ملف.</li>
    <li>ربط جيد بين التحليل الجنائي والواجهة الرسومية وقاعدة البيانات.</li>
    <li>وجود أكثر من نوع تصدير يجعل المشروع مناسبًا للعرض الأكاديمي والعملي.</li>
    <li>استخدام SHA-256 يضيف بعدًا مهمًا في chain of custody.</li>
  </ul>

  <h3>ملاحظات تطويرية مقترحة</h3>
  <ul>
    <li>إضافة دعم أوسع لبعض الصيغ مثل HEIC في استخراج الأبعاد والـ GPS بشكل أكثر اكتمالًا.</li>
    <li>إضافة اختبارات automated tests للوحدات الحساسة مثل decoding وanomaly rules.</li>
    <li>إضافة سجل audit log لكل إجراء ينفذه المستخدم داخل التطبيق.</li>
    <li>توسيع التقرير ليشمل صور مصغرة للأدلة نفسها عند الحاجة.</li>
  </ul>

  <h2>13. الخلاصة</h2>
  <p>
    GeoTrace Forensics مشروع قوي ومنظم لفكرة التحقيق الرقمي المعتمد على صور الأدلة. المشروع يطبق دورة عمل متكاملة تبدأ
    من إدخال الصورة وتوثيق سلامتها، ثم تحليل EXIF وGPS والشذوذات، وتنتهي بإخراج تقارير قابلة للمراجعة. الحالة العملية
    Case {case_info['case_id']} أوضحت أن الأداة قادرة على التقاط معلومات الجهاز ووقت الالتقاط والتحقق من سلامة الملف،
    وفي الوقت نفسه الإشارة بوضوح إلى غياب GPS ووجود تعديل لاحق على الملف داخل نظام التشغيل.
  </p>
</body>
</html>
"""
    REPORT_HTML.write_text(html_body, encoding="utf-8")


def convert_to_docx():
    subprocess.run(
        [
            "textutil",
            "-convert",
            "docx",
            str(REPORT_HTML),
            "-output",
            str(REPORT_DOCX),
        ],
        check=True,
    )


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def set_cell_text(cell, text, bold=False, rtl=True):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if rtl:
        set_rtl(paragraph)
    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18 if level == 1 else 15 if level == 2 else 13)
    run.font.color.rgb = RGBColor(0x12, 0x32, 0x4A)


def add_rtl_paragraph(document, text, bold=False):
    paragraph = document.add_paragraph()
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = bold
    return paragraph


def add_ltr_code_block(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)


def add_figure(document, image_path, caption):
    document.add_picture(str(image_path), width=Cm(16.5))
    caption_paragraph = document.add_paragraph()
    set_rtl(caption_paragraph)
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x5C, 0x67, 0x70)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        shade_cell(header_cells[index], "264653")
        for run in header_cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(row_cells[index], value)

    document.add_paragraph("")


def add_heading_en(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18 if level == 1 else 15 if level == 2 else 13)
    run.font.color.rgb = RGBColor(0x12, 0x32, 0x4A)


def add_paragraph_en(document, text, bold=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = bold
    return paragraph


def add_figure_en(document, image_path, caption):
    document.add_picture(str(image_path), width=Cm(16.5))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x5C, 0x67, 0x70)


def add_table_en(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True, rtl=False)
        shade_cell(header_cells[index], "264653")
        header_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in header_cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(row_cells[index], value, rtl=False)
            row_cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph("")


def build_docx(case_data):
    case_info = case_data["case"]
    image_info = case_data["images"][0]
    summary = case_data["timeline"]["summary"]
    anomaly_rows = image_info.get("anomalies", [])

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    document.add_picture(str(LOGO_PATH), width=Cm(12.5))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("توثيق احترافي لمشروع GeoTrace Forensics")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(0x12, 0x32, 0x4A)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Desktop Digital Forensics Tool for Image Metadata, Geolocation, Timeline, and Reporting")
    subrun.font.name = "Arial"
    subrun.font.size = Pt(10.5)

    add_table(
        document,
        ["البند", "القيمة"],
        [
            ["اسم المشروع", "GeoTrace Forensics"],
            ["نوع المشروع", "Desktop Forensic Investigation Application built with Python and PyQt5"],
            ["تاريخ إعداد الوثيقة", "2026-04-25"],
            [f"الحالة العملية", f"Case #{case_info['case_id']} - {case_info['case_name']}"],
        ],
    )

    add_heading(document, "1. الملخص التنفيذي")
    add_rtl_paragraph(
        document,
        "هذا المشروع عبارة عن أداة تحقيق رقمي مكتبية لتحليل صور الأدلة الرقمية. يقوم النظام باستيراد الصور، واستخراج بيانات EXIF، وفك إحداثيات GPS، والتحقق من سلامة الأدلة عبر SHA-256، ثم بناء خط زمني واكتشاف الشذوذات وتصدير النتائج في أكثر من صيغة.",
    )
    add_rtl_paragraph(
        document,
        "من خلال قراءة الكود والمخرجات الفعلية، يتضح أن المشروع يطبق workflow جنائي متكامل يبدأ من intake evidence وينتهي بتقارير قابلة للمراجعة والتسليم.",
    )

    add_heading(document, "2. هدف المشروع والمشكلة التي يحلها")
    add_rtl_paragraph(
        document,
        "في التحقيقات الرقمية، صور الهواتف والكاميرات قد تحتوي على أدلة مهمة مثل وقت الالتقاط، نوع الجهاز، البرنامج المستخدم في التعديل، والموقع الجغرافي. المشروع يجمع هذه البيانات في بيئة واحدة، ويحولها إلى حزمة أدلة منظمة تساعد على التحليل السريع والدقيق.",
    )
    add_figure(document, ASSETS_DIR / "architecture.png", "شكل 1: التسلسل العام لعمل النظام من إدخال الأدلة حتى إخراج النتائج.")

    add_heading(document, "3. المزايا الأساسية للمشروع")
    add_table(
        document,
        ["الميزة", "الوصف"],
        [
            ["EXIF Extraction", "قراءة بيانات الكاميرا ووقت الالتقاط والبرنامج والوسوم الخام من كل صورة."],
            ["GPS Decoding", "تحويل إحداثيات EXIF إلى قيم decimal يمكن الاعتماد عليها في التحقيق."],
            ["Timeline Analysis", "ترتيب الصور زمنيًا وحساب الفروق الزمنية والمكانية."],
            ["Integrity Verification", "التحقق من سلامة الأدلة باستخدام SHA-256."],
            ["Anomaly Detection", "اكتشاف غياب EXIF وGPS والوقت المشبوه وبرامج التعديل."],
            ["Forensic Outputs", "إنتاج PDF وExcel وJSON وخرائط HTML تفاعلية."],
        ],
    )

    add_heading(document, "4. بنية المشروع والملفات الرئيسية")
    add_rtl_paragraph(
        document,
        "المشروع منظم على شكل واجهة رسومية، وحدات تحليل مستقلة، قاعدة بيانات SQLite، ومجلد إخراج لكل قضية. هذا الفصل يجعل الكود أوضح وأسهل في الصيانة والتطوير.",
    )
    add_figure(document, ASSETS_DIR / "project_tree.png", "شكل 2: لقطة مبنية من هيكل المشروع الفعلي والملفات الأهم داخله.")
    add_table(
        document,
        ["الملف أو الوحدة", "الدور داخل المشروع"],
        [
            ["main.py", "نقطة الدخول الرئيسية للتطبيق وتحميل ملف التنسيق."],
            ["ui/main_window.py", "الواجهة الرسومية الخاصة بإدارة القضايا والأدلة والمخرجات."],
            ["modules/exif_extractor.py", "استخراج البيانات الوصفية من الصور."],
            ["modules/gps_decoder.py", "فك وتحويل وفحص إحداثيات GPS."],
            ["modules/anomaly_detector.py", "تطبيق قواعد اكتشاف الشذوذات الجنائية."],
            ["modules/timeline_generator.py", "إنشاء التسلسل الزمني وربط الصور المتقاربة."],
            ["modules/map_generator.py", "توليد خريطة HTML تفاعلية باستخدام folium."],
            ["modules/report_generator.py", "إنشاء تقرير PDF منسق."],
            ["modules/export_manager.py", "إنتاج ملفات Excel وJSON."],
            ["modules/db_manager.py", "إدارة التخزين في SQLite."],
        ],
    )

    add_heading(document, "5. شرح سير العمل داخل التطبيق")
    for item in [
        "إنشاء قضية جديدة وإدخال اسم القضية واسم المحقق.",
        "استيراد الصور وربطها بالقضية الحالية.",
        "حساب SHA-256 لكل صورة وتخزين المسار والحجم.",
        "استخراج EXIF مثل الكاميرا ووقت الالتقاط والبرنامج المستخدم.",
        "فك إحداثيات GPS عند وجودها.",
        "تطبيق قواعد كشف الشذوذات على كل صورة.",
        "ترتيب النتائج زمنيًا وإنشاء correlations عند وجود تقارب زمني ومكاني.",
        "إخراج النتائج في PDF وExcel وJSON وHTML Map.",
    ]:
        add_rtl_paragraph(document, f"• {item}")

    add_heading(document, "6. شرح الواجهة الرسومية")
    add_rtl_paragraph(
        document,
        "الواجهة الرئيسية مبنية باستخدام PyQt5، وتجمع إدارة القضايا، شريط العمليات، وجدول النتائج مع تبويبات لعرض الميتاداتا، الشذوذات، وRaw EXIF. هذا يجعل المحقق قادرًا على استعراض الأدلة من شاشة واحدة.",
    )
    add_figure(document, ASSETS_DIR / "ui_code.png", "شكل 3: لقطة من كود الواجهة الرئيسية توضح بناء عناصر التحكم ومسار العمل.")

    add_heading(document, "7. قاعدة البيانات والتخزين الجنائي")
    add_rtl_paragraph(
        document,
        "يعتمد المشروع على قاعدة SQLite تتكون من الجداول: cases وimages وmetadata وanomalies. هذا التصميم يحقق فصلًا واضحًا بين بيانات القضية، وبيانات الأدلة، والنتائج التحليلية، والملاحظات الجنائية.",
    )
    add_table(
        document,
        ["الجدول", "الغرض"],
        [
            ["cases", "تخزين اسم القضية واسم المحقق ووقت الإنشاء."],
            ["images", "تسجيل الأدلة الرقمية ومساراتها وقيم الهاش الخاصة بها."],
            ["metadata", "تخزين نتائج EXIF وGPS والأبعاد والبرنامج المستخدم."],
            ["anomalies", "تسجيل كل شذوذ أو ملاحظة تحقيقية مرتبطة بالصورة."],
        ],
    )

    add_heading(document, f"8. التحليل العملي للحالة Case {case_info['case_id']}")
    add_rtl_paragraph(
        document,
        f"تم استخدام ملف التصدير الحقيقي الموجود في output/case_4/case_4_export.json. تحتوي الحالة على صورة واحدة هي {image_info['file_name']} التقطت بواسطة {image_info.get('camera_make') or 'N/A'} {image_info.get('camera_model') or ''}.",
    )
    add_table(
        document,
        ["العنصر", "القيمة"],
        [
            ["Case ID", case_info["case_id"]],
            ["Case Name", case_info["case_name"]],
            ["Investigator", case_info["investigator_name"]],
            ["Image File", image_info["file_name"]],
            ["Capture Time", image_info.get("date_taken") or "Unavailable"],
            ["GPS Status", f"{summary.get('images_with_gps', 0)} image(s) with GPS"],
            ["Integrity Status", image_info.get("integrity_status") or "Unknown"],
            ["Total Anomalies", summary.get("total_anomalies", 0)],
            ["SHA-256", image_info["sha256_hash"]],
        ],
    )
    add_figure(document, ASSETS_DIR / "case4_summary.png", "شكل 4: لقطة ملخصة لأهم نتائج Case 4 من ملف التصدير الفعلي.")
    add_rtl_paragraph(
        document,
        "أظهرت الحالة أن الملف سليم جنائيًا من ناحية الهاش لأن SHA-256 الحالي يطابق القيمة المخزنة. لكن لم يتم العثور على GPS، وبالتالي لا يمكن تحديد الموقع الجغرافي مباشرة. كما اكتشف النظام أن وقت التعديل على نظام الملفات جاء بعد وقت الالتقاط المسجل داخل EXIF، فتم تسجيل شذوذ منخفض الخطورة.",
    )
    add_table(
        document,
        ["نوع الشذوذ", "الخطورة", "الوصف"],
        [[row["anomaly_type"], row["severity"], row["description"]] for row in anomaly_rows],
    )
    add_figure(document, ASSETS_DIR / "json_snapshot.png", "شكل 5: لقطة من ملف JSON المولد فعليًا للحالة.")

    add_heading(document, "9. شرح الكود الأساسي")
    add_rtl_paragraph(
        document,
        "في هذا الجزء تم اختيار أهم الأجزاء البرمجية التي تمثل قلب المشروع: استخراج الميتاداتا، كشف الشذوذات، وبناء الخط الزمني.",
    )

    add_heading(document, "9.1 استخراج EXIF", level=3)
    add_rtl_paragraph(
        document,
        "تعتمد وحدة exif_extractor.py على قراءة أولية بالصورة باستخدام Pillow، ثم محاولة قراءة الوسوم عبر exifread. بعد ذلك يتم دمج النتائج في بنية موحدة حتى تكون جاهزة لباقي الوحدات.",
    )
    add_figure(document, ASSETS_DIR / "exif_code.png", "شكل 6: لقطة من كود استخراج EXIF ودمج الوسوم في Metadata موحدة.")
    add_ltr_code_block(document, snippet("modules/exif_extractor.py", 41, 90))

    add_heading(document, "9.2 كشف الشذوذات", level=3)
    add_rtl_paragraph(
        document,
        "وحدة anomaly_detector.py هي القلب التحليلي للمشروع، لأنها تطبق قواعد جنائية على الصورة مثل غياب EXIF أو GPS، والوقت غير المنطقي، ووجود برامج تعديل، والفارق بين وقت الالتقاط ووقت التعديل على نظام الملفات.",
    )
    add_figure(document, ASSETS_DIR / "anomaly_code.png", "شكل 7: لقطة من كود القواعد الجنائية لاكتشاف المشكلات داخل Metadata.")
    add_ltr_code_block(document, snippet("modules/anomaly_detector.py", 28, 91))

    add_heading(document, "9.3 بناء الخط الزمني", level=3)
    add_rtl_paragraph(
        document,
        "وحدة timeline_generator.py ترتب الصور زمنيًا وتحسب الفروق الزمنية والمسافات بين الصور باستخدام معادلة Haversine. وعند تحقق شروط القرب الزمني والمكاني، يتم إنشاء correlation مفيد للتحقيق.",
    )
    add_figure(document, ASSETS_DIR / "timeline_code.png", "شكل 8: لقطة من كود التحليل الزمني وحساب التقارب المكاني بين الصور.")
    add_ltr_code_block(document, snippet("modules/timeline_generator.py", 33, 118))

    add_heading(document, "10. المخرجات التي ينتجها النظام")
    for item in [
        "PDF Report: تقرير منسق يلخص القضية والأدلة والخط الزمني والشذوذات.",
        "Excel Export: مناسب للفرز والتحليل اليدوي ومراجعة البيانات.",
        "JSON Export: مناسب للتكامل مع أدوات أخرى أو التحليل البرمجي.",
        "HTML Map: خريطة تفاعلية للصور التي تحتوي على GPS.",
    ]:
        add_rtl_paragraph(document, f"• {item}")
    add_rtl_paragraph(
        document,
        f"ملاحظة: الحالة العملية Case {case_info['case_id']} لا تحتوي على GPS، لذلك لا توجد خريطة خاصة بها ضمن المخرجات الحالية، رغم أن الميزة موجودة في الكود ومفعلة للمشاهد التي تحتوي على إحداثيات.",
    )

    add_heading(document, "11. طريقة تشغيل المشروع")
    add_ltr_code_block(document, "pip install -r requirements.txt\npython main.py")
    add_rtl_paragraph(
        document,
        "يعتمد المشروع على المكتبات: PyQt5, Pillow, exifread, folium, reportlab, openpyxl, pandas. بعد تشغيل التطبيق تظهر الواجهة مباشرة ويمكن بدء إنشاء القضايا وإضافة الصور وتحليلها.",
    )

    add_heading(document, "12. نقاط القوة والملاحظات التطويرية")
    add_rtl_paragraph(document, "نقاط القوة:")
    for item in [
        "تقسيم واضح للوحدات البرمجية وسهولة فهم مسؤولية كل ملف.",
        "ربط جيد بين التحليل الجنائي والواجهة الرسومية وقاعدة البيانات.",
        "وجود أكثر من صيغة تصدير يزيد من قيمة المشروع عمليًا وأكاديميًا.",
        "استخدام SHA-256 يدعم chain of custody بشكل مهم.",
    ]:
        add_rtl_paragraph(document, f"• {item}")
    add_rtl_paragraph(document, "ملاحظات تطويرية مقترحة:")
    for item in [
        "إضافة دعم أوسع لصيغ مثل HEIC في استخراج الأبعاد وGPS.",
        "إضافة اختبارات automated tests للوحدات التحليلية المهمة.",
        "إضافة audit log واضح لكل إجراء ينفذه المستخدم.",
        "توسيع التقارير لتضمين صور مصغرة للأدلة نفسها.",
    ]:
        add_rtl_paragraph(document, f"• {item}")

    add_heading(document, "13. الخلاصة")
    add_rtl_paragraph(
        document,
        f"GeoTrace Forensics مشروع منظم وقوي في مجال التحقيق الرقمي المعتمد على صور الأدلة. أوضحت الحالة العملية Case {case_info['case_id']} أن النظام قادر على استخراج معلومات الجهاز ووقت الالتقاط والتحقق من سلامة الملف، وفي الوقت نفسه التنبيه إلى غياب GPS ووجود تعديل لاحق على مستوى نظام الملفات. لذلك يمكن اعتباره مشروعًا أكاديميًا احترافيًا وقابلًا للتطوير إلى أداة أوسع نطاقًا.",
    )

    document.save(REPORT_DOCX)


def build_docx_english(case_data):
    case_info = case_data["case"]
    image_info = case_data["images"][0]
    summary = case_data["timeline"]["summary"]
    anomaly_rows = image_info.get("anomalies", [])

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    document.add_picture(str(LOGO_PATH), width=Cm(12.5))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Professional Project Documentation for GeoTrace Forensics")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(0x12, 0x32, 0x4A)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Desktop Digital Forensics Tool for Image Metadata, Geolocation, Timeline, and Reporting")
    subrun.font.name = "Arial"
    subrun.font.size = Pt(10.5)

    add_table_en(
        document,
        ["Item", "Value"],
        [
            ["Project Name", "GeoTrace Forensics"],
            ["Project Type", "Desktop Forensic Investigation Application built with Python and PyQt5"],
            ["Document Date", "2026-04-25"],
            ["Case Used in This Documentation", f"Case #{case_info['case_id']} - {case_info['case_name']}"],
        ],
    )

    add_heading_en(document, "1. Executive Summary")
    add_paragraph_en(
        document,
        "GeoTrace Forensics is a desktop digital forensics application designed to analyze image evidence in a structured and investigator-friendly way. The system imports images, extracts EXIF metadata, decodes GPS coordinates when available, verifies file integrity using SHA-256, builds a chronological timeline, detects forensic anomalies, and exports the findings in multiple formats.",
    )
    add_paragraph_en(
        document,
        "After reviewing the actual codebase and generated case outputs, the project clearly implements a complete forensic workflow rather than a simple metadata viewer. It starts with evidence intake and ends with usable reporting artifacts for review and submission.",
    )

    add_heading_en(document, "2. Project Goal and Problem Statement")
    add_paragraph_en(
        document,
        "In digital investigations, image files can contain highly valuable evidence such as capture time, device identity, editing software, and geolocation. The challenge is that this information is usually scattered across metadata fields and is not easy to analyze manually at scale. GeoTrace Forensics solves this by collecting these details into one organized environment and turning them into a structured evidence package.",
    )
    add_figure_en(document, ASSETS_DIR / "architecture.png", "Figure 1. High-level workflow from evidence intake to final forensic outputs.")

    add_heading_en(document, "3. Core Project Features")
    add_table_en(
        document,
        ["Feature", "Description"],
        [
            ["EXIF Extraction", "Reads camera data, timestamps, software fields, and raw metadata from each image."],
            ["GPS Decoding", "Converts embedded GPS data into investigator-friendly decimal coordinates."],
            ["Timeline Analysis", "Sorts images chronologically and measures time and distance gaps."],
            ["Integrity Verification", "Uses SHA-256 hashing to confirm evidence integrity."],
            ["Anomaly Detection", "Flags missing EXIF, missing GPS, suspicious timestamps, and editing traces."],
            ["Forensic Outputs", "Generates PDF, Excel, JSON, and interactive HTML map artifacts."],
        ],
    )

    add_heading_en(document, "4. Project Structure and Main Files")
    add_paragraph_en(
        document,
        "The project is organized into a clear separation between the desktop interface, forensic analysis modules, the SQLite storage layer, and case output folders. This design improves readability, maintenance, and extensibility.",
    )
    add_figure_en(document, ASSETS_DIR / "project_tree.png", "Figure 2. Snapshot of the real project structure and the most important files.")
    add_table_en(
        document,
        ["File or Module", "Responsibility"],
        [
            ["main.py", "Application entry point and stylesheet loading."],
            ["ui/main_window.py", "Desktop UI for case management, evidence intake, and artifact generation."],
            ["modules/exif_extractor.py", "Extracts image metadata from EXIF sources."],
            ["modules/gps_decoder.py", "Decodes, validates, and formats GPS coordinates."],
            ["modules/anomaly_detector.py", "Applies forensic rules to identify suspicious metadata conditions."],
            ["modules/timeline_generator.py", "Builds chronological ordering and correlation analysis."],
            ["modules/map_generator.py", "Creates interactive HTML maps using folium."],
            ["modules/report_generator.py", "Builds the PDF forensic report."],
            ["modules/export_manager.py", "Generates Excel and JSON exports."],
            ["modules/db_manager.py", "Handles SQLite persistence and retrieval."],
        ],
    )

    add_heading_en(document, "5. Application Workflow")
    for item in [
        "Create a new case and enter the case name and investigator name.",
        "Import one or more image files into the active case.",
        "Compute SHA-256 for each image and store the file path and size.",
        "Extract EXIF data such as camera, timestamp, and software fields.",
        "Decode GPS coordinates when they are available in the metadata.",
        "Apply anomaly-detection rules to each imported image.",
        "Build a timeline and create correlations for closely related events.",
        "Export the results as PDF, Excel, JSON, and interactive HTML map artifacts.",
    ]:
        add_paragraph_en(document, f"• {item}")

    add_heading_en(document, "6. User Interface Overview")
    add_paragraph_en(
        document,
        "The main interface is built with PyQt5 and combines case management, action controls, a results table, and detailed review tabs. The investigator can create or load a case, import images, generate reports, verify integrity, and inspect metadata, anomalies, and raw EXIF from one screen.",
    )
    add_figure_en(document, ASSETS_DIR / "ui_code.png", "Figure 3. Snapshot from the main UI controller showing the structure of the workflow.")

    add_heading_en(document, "7. Database Design and Forensic Storage")
    add_paragraph_en(
        document,
        "The project uses a local SQLite database with four main tables: cases, images, metadata, and anomalies. This schema cleanly separates case information, evidence intake, extracted metadata, and forensic findings, which is a strong design choice for a project of this type.",
    )
    add_table_en(
        document,
        ["Table", "Purpose"],
        [
            ["cases", "Stores the case name, investigator name, and creation time."],
            ["images", "Stores evidence files, file paths, and cryptographic hashes."],
            ["metadata", "Stores EXIF results, GPS values, dimensions, and software metadata."],
            ["anomalies", "Stores every anomaly or forensic finding linked to an image."],
        ],
    )

    add_heading_en(document, f"8. Practical Analysis of Case {case_info['case_id']}")
    add_paragraph_en(
        document,
        f"This documentation uses the real export file located at output/case_4/case_4_export.json. The case contains one image named {image_info['file_name']} captured by {image_info.get('camera_make') or 'N/A'} {image_info.get('camera_model') or ''}.",
    )
    add_table_en(
        document,
        ["Item", "Value"],
        [
            ["Case ID", case_info["case_id"]],
            ["Case Name", case_info["case_name"]],
            ["Investigator", case_info["investigator_name"]],
            ["Image File", image_info["file_name"]],
            ["Capture Time", image_info.get("date_taken") or "Unavailable"],
            ["GPS Status", f"{summary.get('images_with_gps', 0)} image(s) with GPS"],
            ["Integrity Status", image_info.get("integrity_status") or "Unknown"],
            ["Total Anomalies", summary.get("total_anomalies", 0)],
            ["SHA-256", image_info["sha256_hash"]],
        ],
    )
    add_figure_en(document, ASSETS_DIR / "case4_summary.png", "Figure 4. Snapshot summarizing the key findings from Case 4.")
    add_paragraph_en(
        document,
        "The case shows that the file passed integrity verification because the current SHA-256 hash matched the stored hash recorded during evidence intake. However, no GPS coordinates were found, which means the image cannot be geolocated directly. The system also detected that the filesystem modification time was later than the recorded capture time, so it flagged a low-severity anomaly.",
    )
    add_table_en(
        document,
        ["Anomaly Type", "Severity", "Description"],
        [[row["anomaly_type"], row["severity"], row["description"]] for row in anomaly_rows],
    )
    add_figure_en(document, ASSETS_DIR / "json_snapshot.png", "Figure 5. Snapshot from the generated JSON export for the real case output.")

    add_heading_en(document, "9. Explanation of the Key Code")
    add_paragraph_en(
        document,
        "This section focuses on the three most important technical parts of the project: metadata extraction, anomaly detection, and timeline generation.",
    )

    add_heading_en(document, "9.1 EXIF Extraction", level=3)
    add_paragraph_en(
        document,
        "The exif_extractor.py module reads general image information using Pillow and then attempts detailed EXIF extraction using exifread. The two sources are merged into one normalized metadata object, which makes the pipeline more resilient across different image formats.",
    )
    add_figure_en(document, ASSETS_DIR / "exif_code.png", "Figure 6. Snapshot from the EXIF extraction logic and metadata normalization flow.")
    add_ltr_code_block(document, snippet("modules/exif_extractor.py", 41, 90))

    add_heading_en(document, "9.2 Anomaly Detection", level=3)
    add_paragraph_en(
        document,
        "The anomaly_detector.py module is the analytical core of the project. It applies forensic rules that check for missing EXIF, missing GPS, unsupported timestamps, editing software traces, and differences between capture time and filesystem modification time.",
    )
    add_figure_en(document, ASSETS_DIR / "anomaly_code.png", "Figure 7. Snapshot from the forensic rules used to detect suspicious metadata conditions.")
    add_ltr_code_block(document, snippet("modules/anomaly_detector.py", 28, 91))

    add_heading_en(document, "9.3 Timeline Generation", level=3)
    add_paragraph_en(
        document,
        "The timeline_generator.py module sorts images chronologically and calculates both time gaps and geographic distance using the Haversine formula when coordinates exist. If events are sufficiently close in time and space, the system creates a correlation entry to support the investigation narrative.",
    )
    add_figure_en(document, ASSETS_DIR / "timeline_code.png", "Figure 8. Snapshot from the code that builds the timeline and spatial correlation logic.")
    add_ltr_code_block(document, snippet("modules/timeline_generator.py", 33, 118))

    add_heading_en(document, "10. System Outputs")
    for item in [
        "PDF Report: A formatted forensic summary of the case, evidence, timeline, and anomalies.",
        "Excel Export: A structured format suitable for filtering, sorting, and manual review.",
        "JSON Export: A machine-readable format suitable for later automation or integration.",
        "HTML Map: An interactive map for cases that contain GPS-enabled images.",
    ]:
        add_paragraph_en(document, f"• {item}")
    add_paragraph_en(
        document,
        f"Note: Case {case_info['case_id']} does not contain GPS coordinates, so there is no case-specific map output for this particular example, even though the mapping feature is implemented in the project.",
    )

    add_heading_en(document, "11. How to Run the Project")
    add_ltr_code_block(document, "pip install -r requirements.txt\npython main.py")
    add_paragraph_en(
        document,
        "The project depends on PyQt5, Pillow, exifread, folium, reportlab, openpyxl, and pandas. After launching the application, the investigator can create a case, import images, review findings, and generate exports from the GUI.",
    )

    add_heading_en(document, "12. Strengths and Suggested Improvements")
    add_paragraph_en(document, "Strengths:", bold=True)
    for item in [
        "Clear modular structure with well-separated responsibilities.",
        "Good integration between forensic analysis, the GUI, and the database layer.",
        "Multiple export formats increase the practical value of the project.",
        "SHA-256 hashing supports evidence integrity and chain-of-custody practices.",
    ]:
        add_paragraph_en(document, f"• {item}")
    add_paragraph_en(document, "Suggested Improvements:", bold=True)
    for item in [
        "Expand support for image formats such as HEIC, especially for dimensions and GPS extraction.",
        "Add automated tests for decoding and anomaly-detection logic.",
        "Introduce a clearer audit log for every investigator action.",
        "Extend reports to include evidence thumbnails when needed.",
    ]:
        add_paragraph_en(document, f"• {item}")

    add_heading_en(document, "13. Conclusion")
    add_paragraph_en(
        document,
        f"GeoTrace Forensics is a strong and well-structured academic project in the area of image-based digital forensics. The practical example from Case {case_info['case_id']} shows that the tool can extract device information, record the capture time, verify file integrity, and clearly report missing GPS data and post-capture modification indicators. Overall, the project is professional in structure and can be extended into a more advanced forensic analysis platform.",
    )

    document.save(REPORT_DOCX_EN)


def main():
    ensure_dirs()
    case_data = load_json(CASE_JSON_PATH)
    create_assets(case_data)
    build_html(case_data)
    build_docx(case_data)
    build_docx_english(case_data)
    print(f"HTML report: {REPORT_HTML}")
    print(f"DOCX report: {REPORT_DOCX}")
    print(f"English DOCX report: {REPORT_DOCX_EN}")


if __name__ == "__main__":
    main()
